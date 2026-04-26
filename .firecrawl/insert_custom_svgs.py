"""Insert the five originally-authored SVGs (rendered to PNG) into the chapter
docx files above their FIGURE DESCRIPTION tables. Self-attributed since these
are original work for this OER, CC BY 4.0.
"""
from pathlib import Path
import shutil
import sys

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
BACKUP_DIR = ROOT / ".firecrawl" / "backups"

CHAPTER_FILES = {
    "02": "Chapter_02_Unit_Systems_and_Dimensional_Analysis.docx",
    "07": "Chapter_07_Chemical_Nomenclature.docx",
    "08": "Chapter_08_Mole_Concept_and_Chemical_Formulas.docx",
    "09": "Chapter_09_Chemical_Calculations_and_Equations.docx",
    "10": "Chapter_10_States_of_Matter.docx",
}

SELF_ATTRIB = (
    "Original figure for the CHEM 139 OER textbook (2026 ed.). "
    "Licensed CC BY 4.0."
)

PLAN = {
    "02": [("Figure 2.1 — Two ways to think about density.",
            "fig2-1_density_cubes.png")],
    "07": [("Figure 7.1 — Naming decision tree.",
            "fig7-1_naming_tree.png")],
    "08": [("Figure 8.1 — The mole map (chemical-amount triangle).",
            "fig8-1_mole_map.png")],
    "09": [("Figure 9.1 — Stoichiometry road map.",
            "fig9-1_stoichiometry_roadmap.png")],
    "10": [("Figure 10.1 — Hierarchy of intermolecular forces.",
            "fig10-1_imf_hierarchy.png")],
}


def find_figure_table(doc, key_prefix):
    for t in doc.tables:
        try:
            first_text = t.rows[0].cells[0].text.strip()
        except IndexError:
            continue
        if first_text.startswith(key_prefix):
            return t
    return None


def already_has_image_above(table):
    """Return True if the paragraph immediately preceding the table already
    contains a drawing (i.e. an image was inserted earlier — avoid duplication)."""
    prev = table._element.getprevious()
    # Walk backward across attribution-text paragraphs to the image paragraph.
    for _ in range(3):
        if prev is None:
            return False
        if prev.tag.endswith('}p'):
            if prev.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None:
                return True
        prev = prev.getprevious()
    return False


def insert_image_block_before_table(doc, table, image_path, attribution, width_in=4.8):
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(image_path), width=Inches(width_in))

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(attribution)
    cap_run.italic = True
    cap_run.font.size = Pt(9)

    img_el = img_p._element
    cap_el = cap_p._element
    img_el.getparent().remove(img_el)
    cap_el.getparent().remove(cap_el)
    table._element.addprevious(img_el)
    table._element.addprevious(cap_el)


def process_chapter(num, filename, items):
    src = ROOT / filename
    if not src.exists():
        sys.exit(f"Missing source: {src}")
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{src.stem}.before-custom-svgs.docx"
    shutil.copy2(src, backup)

    doc = Document(str(src))
    inserted = []
    skipped = []
    for header, img_name in items:
        img_path = ROOT / "Images" / f"Chapter_{num}" / "embedded" / img_name
        if not img_path.exists():
            sys.exit(f"Missing image: {img_path}")
        t = find_figure_table(doc, header)
        if t is None:
            sys.exit(f"  Ch {num}: could not find table for {header!r}")
        if already_has_image_above(t):
            skipped.append(header)
            continue
        insert_image_block_before_table(doc, t, img_path, SELF_ATTRIB)
        inserted.append(header)

    if inserted:
        doc.save(str(src))
    print(f"Ch {num}: inserted {len(inserted)} new image(s)" +
          (f", skipped {len(skipped)} (already present)" if skipped else ""))
    for h in inserted:
        print(f"   + {h}")
    for h in skipped:
        print(f"   . {h} (skipped)")


def main():
    for num, items in PLAN.items():
        filename = CHAPTER_FILES[num]
        process_chapter(num, filename, items)


if __name__ == "__main__":
    main()
