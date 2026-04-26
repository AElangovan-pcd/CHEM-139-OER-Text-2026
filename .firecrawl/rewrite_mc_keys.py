"""Append a worked-rationale block to each chapter's MC Practice Test answer key.

For every chapter that has a "Multiple Choice Practice Test", this finds
the "Answer Key:" paragraph and inserts, immediately after it, an
"Answer Key with Worked Rationales" sub-section: one paragraph per
question giving the correct letter plus a brief explanation. Math
questions get a one-line factor-label calculation; conceptual questions
get a single sentence on why the chosen answer is right.

The terse one-line "Answer Key: 1-B 2-C ..." summary is preserved as-is
above the rationale block, so quick-reference scanning still works.

Idempotent: skips a chapter if "Worked Rationales" already appears.
"""
from copy import deepcopy
from pathlib import Path
import shutil
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
BACKUP_DIR = ROOT / ".firecrawl" / "backups"

W_PPR = qn("w:pPr")


def _add_run(p, text, *, bold=False, italic=False):
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def _insert_paragraph_after(anchor):
    new_p = OxmlElement("w:p")
    src_pPr = anchor._element.find(W_PPR)
    if src_pPr is not None:
        new_p.append(deepcopy(src_pPr))
    anchor._element.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


# Each value is a list of (qno_int, letter, rationale_text) for that chapter.
RATIONALES = {
    "Chapter_01_Science_and_Measurement.docx": [
        (1, "B", "Chemistry is the study of matter and the changes it undergoes (standard definition)."),
        (2, "B", "Atomic theory — atoms make up matter and combine in fixed ratios — is a well-tested broad explanation: a scientific theory."),
        (3, "B", "E = mc² states a quantitative pattern (energy and mass are equivalent at rest) — a scientific law, summarising behaviour without explaining mechanism."),
        (4, "C", "A hypothesis must be falsifiable: it has to make a risky prediction that observation could in principle refute."),
        (5, "B", "Ethical questions are normative (\"is X right?\") rather than empirical; they cannot be settled by measurement."),
        (6, "B", "Defined relationships are exact. \"60 seconds per 1 minute\" comes from the SI definition; the others involve measurement."),
        (7, "C", "The data cluster tightly (precise) but sit far from the true value 7.00 — accurate they are not."),
        (8, "B", "0.005020 has 4 sig figs: leading zeros don't count; the captive 0 between 5 and 2 counts; the trailing 0 (after a decimal point) counts."),
        (9, "C", "4.0080 has 5 sig figs: 4, the captive 0, 0, 8, and the trailing 0 (made significant by the decimal point)."),
        (10, "C", "Addition is limited by the least precise decimal place; \"12 g\" has none, so 3.14 + 0.5 + 12 = 15.64 rounds to 16 g."),
        (11, "B", "(5.0 × 10²) × (2.50 × 10⁻⁴) = 1.25 × 10⁻¹; multiplication is limited by the input with fewest sig figs (5.0 has 2) → 1.3 × 10⁻¹."),
        (12, "B", "3 sig figs from \"3,200,000\" requires explicit precision: 3.20 × 10⁶."),
        (13, "A", "0.0000102 → move the decimal 5 places to the right → 1.02 × 10⁻⁵."),
        (14, "A", "Random errors scatter symmetrically around the true value; averaging cancels them. Systematic errors all bias the same direction and survive averaging."),
        (15, "B", "Systematic error is a consistent bias (a miscalibrated balance, an offset thermometer); the cure is recalibration, not more measurements."),
        (16, "C", "(8.40 × 10⁵) / (2.0 × 10²) = 4.2 × 10³; 2 sig figs are the limit (from 2.0)."),
        (17, "A", "4.50 × 10⁻³ − 0.12 × 10⁻³ = 4.38 × 10⁻³; the input \"1.2 × 10⁻⁴\" carries only 2 sig figs, governing the result → 4.4 × 10⁻³."),
        (18, "B", "Kinetic-molecular theory is a broad, well-tested explanation of gas behaviour — a scientific theory."),
        (19, "B", "Forming a tentative explanation that can be tested is the hypothesis step."),
        (20, "B", "1 in = 2.54 cm was defined exactly by international agreement in 1959; the others are measured (or differ slightly across definitions)."),
    ],

    "Chapter_02_Unit_Systems_and_Dimensional_Analysis.docx": [
        (1, "B", "The kilogram (kg) is the SI base unit of mass; the gram is a derived unit."),
        (2, "B", "\"Micro-\" means 10⁻⁶ (e.g., 1 µg = 10⁻⁶ g)."),
        (3, "A", "1 mL ≡ 1 cm³ exactly (the cubic centimetre is the same volume as the millilitre by definition)."),
        (4, "B", "250 cm × (1 m / 100 cm) = 2.50 m; \"cm\" cancels."),
        (5, "B", "V = m / ρ = 84.0 g × (1 mL / 4.20 g) = 20.0 mL; \"g\" cancels."),
        (6, "C", "An object floats when its density is less than that of the surrounding fluid (Archimedes)."),
        (7, "D", "T_K = T_C + 273.15 = 100 + 273.15 = 373.15 K → 373 K."),
        (8, "C", "T_F = T_C × 9/5 + 32 = 37 × 9/5 + 32 = 66.6 + 32 = 98.6 °F."),
        (9, "A", "0.500 lb × (453.59 g / 1 lb) = 226.8 g → 227 g; \"lb\" cancels."),
        (10, "C", "Mass of sugar = 480 g × (12.5 g sugar / 100 g sample) = 60.0 g; \"g sample\" cancels via the percent factor."),
    ],

    "Chapter_03_Basic_Concepts_About_Matter.docx": [
        (1, "C", "Reactivity (with HCl, with O₂, etc.) is a chemical property; density and melting point are physical."),
        (2, "B", "Mass scales with how much sample you have — extensive. Density, MP, BP are intensive."),
        (3, "B", "Boiling water changes phase but the substance (H₂O) is unchanged — physical change."),
        (4, "C", "Granite contains visibly distinct mineral grains (quartz, feldspar, mica) — heterogeneous."),
        (5, "C", "H₂O is a compound (two elements combined chemically). O₂ is elemental; Au is an element; brass is an alloy (mixture)."),
        (6, "C", "Sodium = Na (from Latin natrium)."),
        (7, "C", "Pb = lead (from Latin plumbum)."),
        (8, "B", "Oxygen ≈ 46% of the crust by mass — locked into silicates and oxides."),
        (9, "C", "Oxygen ≈ 65% of the human body by mass — mostly as the O in water."),
        (10, "C", "Distillation separates a homogeneous liquid mixture by exploiting differences in boiling point."),
    ],

    "Chapter_04_Atoms_Molecules_Subatomic_Particles.docx": [
        (1, "B", "Dalton's postulate that all atoms of a given element are identical was modified once isotopes were discovered: same Z, different N."),
        (2, "C", "Electron mass ≈ 1/1836 of a proton — by far the smallest of the three subatomic particles."),
        (3, "B", "The neutron carries no net charge; protons are +1, electrons are −1."),
        (4, "C", "Atoms with the same Z but different N are isotopes (same element, different mass)."),
        (5, "B", "²³⁸U has Z = 92, A = 238 → N = A − Z = 146."),
        (6, "C", "Mass number A = protons + neutrons = 11 + 12 = 23 (²³Na)."),
        (7, "B", "Cl₂ is one of the seven diatomic elements (H₂, N₂, O₂, F₂, Cl₂, Br₂, I₂). He, Cu, Ne are not diatomic."),
        (8, "B", "Ca(NO₃)₂: 1 Ca + 2(N + 3 O) = 1 + 2 + 6 = 9 atoms."),
        (9, "B", "Atomic mass = 0.600(50.0) + 0.400(52.0) = 30.0 + 20.8 = 50.8 u."),
        (10, "B", "Rutherford's gold-foil scattering showed the positive charge and almost all the mass concentrated in a tiny dense nucleus."),
    ],

    "Chapter_05_Electronic_Structure_and_Periodicity.docx": [
        (1, "C", "The periodic law: properties recur as a function of atomic number (Z), not mass — Mendeleev's original mass ordering had several anomalies that Z fixed."),
        (2, "C", "A d subshell holds 5 orbitals (ml = −2, −1, 0, +1, +2)."),
        (3, "C", "Maximum capacity of shell n is 2n² = 2(3)² = 18."),
        (4, "B", "P (Z = 15): 1s² 2s² 2p⁶ 3s² 3p³ — fill the 3p with 3 electrons."),
        (5, "C", "Br (Z = 35) past Ar (Z = 18) fills 4s² 3d¹⁰ 4p⁵ → [Ar] 4s² 3d¹⁰ 4p⁵."),
        (6, "B", "[Kr] 5s² 4d¹⁰ 5p⁴ has 36 + 16 = 52 electrons → Te (tellurium)."),
        (7, "B", "Hund's rule: place one electron in each orbital of a degenerate set with parallel spins before pairing."),
        (8, "D", "Rb is in period 5; atomic radius increases down a group, so Rb > K > Na > Li."),
        (9, "D", "Metallic character increases down a group and decreases across a period; Cs is bottom-left of the listed options → most metallic."),
        (10, "C", "Period 4, Group 12, d-block: Zn (Z = 30, [Ar] 4s² 3d¹⁰)."),
    ],

    "Chapter_06_Chemical_Bonds.docx": [
        (1, "C", "Ionic bonds form by electron transfer between a metal (loses electrons) and a nonmetal (gains them)."),
        (2, "C", "Sulfur is in Group 16 → 6 valence electrons (3s² 3p⁴)."),
        (3, "C", "Mg is in Group 2; it loses two electrons to reach the [Ne] noble-gas configuration → Mg²⁺."),
        (4, "C", "Al³⁺ + 3 Cl⁻ → AlCl₃ (charge balance: +3 = 3 × −1)."),
        (5, "B", "(NH₄)₃PO₄: three +1 ammonium cations balance one −3 phosphate anion."),
        (6, "C", "Ca₃(PO₄)₂: 3 Ca + 2(P + 4 O) = 3 + 2 + 8 = 13 atoms."),
        (7, "C", "NH₃ has 4 electron groups (3 bonds + 1 LP) → tetrahedral electron geometry, trigonal-pyramidal molecular shape."),
        (8, "C", "CH₄ is tetrahedral with bond angle ≈ 109.5° (the regular tetrahedron angle)."),
        (9, "D", "C–F has the largest electronegativity difference of the listed bonds (ΔEN = 4.0 − 2.5 = 1.5) → most polar."),
        (10, "D", "Linear CO₂ has two polar C=O bonds whose dipoles point in opposite directions and cancel → nonpolar overall."),
    ],

    "Chapter_07_Chemical_Nomenclature.docx": [
        (1, "C", "Group-2 cations have only one common charge (+2), so no Roman numeral: simply \"calcium chloride\"."),
        (2, "C", "Iron has multiple oxidation states; in Fe₂O₃ each Fe is +3 → iron(III) oxide."),
        (3, "C", "Phosphate is PO₄³⁻; balance with 3 Na⁺ → Na₃PO₄."),
        (4, "B", "Sulfate is SO₄²⁻; balance with 2 NH₄⁺ → (NH₄)₂SO₄."),
        (5, "B", "Two binary nonmetals: Greek prefixes for both → di-nitrogen pent-oxide → dinitrogen pentoxide."),
        (6, "B", "Binary covalent: \"mono-\" is dropped from the first element → carbon tetrachloride."),
        (7, "C", "Aqueous binary acid (H + halide): \"hydro-\" + halide-stem-\"ic\" + \"acid\" → hydrobromic acid."),
        (8, "B", "Oxoacid from \"sulfite\" (SO₃²⁻): drop \"-ite\", add \"-ous acid\" → sulfurous acid."),
        (9, "D", "\"Chlorate\" is ClO₃⁻; the corresponding acid is HClO₃ (chloric acid)."),
        (10, "A", "Iron(II) has charge +2; sulfide is S²⁻; balance gives 1:1 → FeS."),
    ],

    "Chapter_08_Mole_Concept_and_Chemical_Formulas.docx": [
        (1, "B", "Avogadro's number, NA = 6.022 × 10²³ /mol — the count of entities in 1 mole."),
        (2, "B", "M(NH₃) = 14.01 + 3(1.008) = 17.0 g/mol."),
        (3, "C", "2.00 mol × (18.015 g / 1 mol) = 36.0 g; \"mol\" cancels."),
        (4, "B", "1 mol of any substance contains NA = 6.022 × 10²³ entities."),
        (5, "A", "C₂H₄O₂ → divide subscripts by 2 → CH₂O."),
        (6, "D", "n = 56.1 / 14.027 = 4 → multiply CH₂ subscripts by 4 → C₄H₈."),
        (7, "C", "% O = 2 × 16.00 / 44.01 × 100% = 72.7%."),
        (8, "B", "n = 4.00 g × (1 mol / 4.003 g) = 1.00 mol; \"g\" cancels."),
        (9, "A", "Per 100 g: C 40.0/12.01 = 3.33; H 6.7/1.008 = 6.65; O 53.3/16.00 = 3.33. Divide by 3.33 → 1 : 2 : 1 → CH₂O."),
        (10, "A", "n(H₂O) = 9.0 g × (1 mol / 18.0 g) = 0.500 mol; molecules = 0.500 × 6.022 × 10²³ = 3.0 × 10²³."),
    ],

    "Chapter_09_Chemical_Calculations_and_Equations.docx": [
        (1, "B", "Conservation of mass requires the SAME number of atoms of each element on both sides — that's what balancing enforces. Total moles can change."),
        (2, "C", "N₂ + 3 H₂ → 2 NH₃ balances atoms (2 N, 6 H on each side) with the smallest whole-number coefficients."),
        (3, "B", "One reactant breaking into two products → decomposition (AB → A + B)."),
        (4, "D", "Two compounds swap partners (Ag goes with Cl, Na goes with NO₃) → double replacement (AB + CD → AD + CB)."),
        (5, "C", "Coefficient ratio H₂O : O₂ = 2 : 1, so the mole ratio is 2/1."),
        (6, "C", "n(CO₂) = 1.00 mol C₃H₈ × (3 mol CO₂ / 1 mol C₃H₈) = 3.00 mol; mass = 3.00 mol × 44.01 g/mol = 132 g."),
        (7, "B", "From 2 mol N₂ we'd need 6 mol H₂; only 3 mol H₂ are available → H₂ runs out first → H₂ is limiting."),
        (8, "B", "% yield = 75 / 100 × 100% = 75%."),
        (9, "C", "Cu sits below H in the activity series → Cu cannot displace H₂ from acids → no reaction with HCl."),
        (10, "D", "(aq) means aqueous — dissolved in water."),
    ],

    "Chapter_10_States_of_Matter.docx": [
        (1, "C", "Average kinetic energy of particles is directly proportional to absolute temperature; nothing else listed has that direct dependence."),
        (2, "B", "Glass has no long-range order — it's an amorphous solid (a frozen liquid). NaCl, diamond, ice are crystalline."),
        (3, "C", "Water has O–H bonds, so it can hydrogen-bond. Hydrogen bonding dominates over its weaker dipole–dipole and London components."),
        (4, "D", "All four are nonpolar (London-only). Boiling point tracks polarisability, which scales with electron count: Xe (54 e⁻) > Cl₂ (34) > Ar (18) > He (2)."),
        (5, "C", "Hydrogen bonding requires H bonded to F, O, or N. NH₃ qualifies (N–H); CH₄ and CCl₄ don't (no H on F/O/N), and CO₂ has no H at all."),
        (6, "C", "Higher T means more molecules have enough KE to escape the liquid → vapour pressure rises with temperature."),
        (7, "B", "Boiling occurs when vapour pressure equals the external (atmospheric) pressure; the \"normal\" boiling point uses 1.000 atm."),
        (8, "C", "Specific heat of liquid water is 4.184 J/g·°C — unusually high because of hydrogen bonding."),
        (9, "A", "Water's anomaly: liquid (1.000 g/mL at 4 °C) is denser than solid ice (0.917 g/mL) — which is why ice floats."),
        (10, "B", "q = m × c × ΔT = 50.0 g × 4.184 J/(g·°C) × 10.0 °C = 2092 J ≈ 2090 J."),
    ],
}


def insert_rationales(filename, items):
    src = ROOT / filename
    if not src.exists():
        print(f"Missing: {src}")
        return 0

    backup = BACKUP_DIR / f"{src.stem}.before-mc-rationales.docx"
    if not backup.exists():
        BACKUP_DIR.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, backup)

    doc = Document(str(src))

    # Find the "Answer Key:" paragraph at top level. (None of the chapters
    # put their MC answer key inside a table.)
    target = None
    for p in doc.paragraphs:
        if (p.text or "").strip().startswith("Answer Key:"):
            target = p
            break
    if target is None:
        print(f"WARN: 'Answer Key:' not found in {filename}")
        return 0

    # Idempotency: skip if rationale block is already present.
    nxt = target._element.getnext()
    while nxt is not None:
        nxt_text = "".join(t.text or "" for t in nxt.iter(qn("w:t")))
        if "Worked Rationales" in nxt_text:
            return 0
        nxt = nxt.getnext()

    # Insert the heading paragraph after the terse answer-key line.
    heading = _insert_paragraph_after(target)
    _add_run(heading, "Answer Key — Worked Rationales", bold=True)

    anchor = heading
    for qno, letter, rationale in items:
        p = _insert_paragraph_after(anchor)
        _add_run(p, f"{qno}. ", bold=True)
        _add_run(p, f"{letter}. ")
        _add_run(p, rationale, italic=True)
        anchor = p

    doc.save(str(src))
    return len(items)


def main():
    total = 0
    for filename, items in RATIONALES.items():
        n = insert_rationales(filename, items)
        print(f"{filename}: inserted {n} rationale paragraphs.")
        total += n
    print(f"\nTotal: {total} rationale paragraphs across {len(RATIONALES)} chapters.")


if __name__ == "__main__":
    main()
