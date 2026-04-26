"""Insert sourced OER images above the FIGURE DESCRIPTION tables in Chapter 1.

The figure-description blocks are implemented as Word tables (shaded boxes).
We insert a centered picture paragraph and a 9pt italic attribution paragraph
immediately before each table so the visible image precedes the description,
which is preserved as accessible alt-text per project convention.
"""
from copy import deepcopy
from pathlib import Path
import shutil
import sys

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
SRC = ROOT / "Chapter_01_Science_and_Measurement.docx"
BACKUP = ROOT / ".firecrawl" / "Chapter_01_Science_and_Measurement.before-images.docx"

FIG1_IMG = ROOT / "Images" / "Chapter_01" / "embedded" / "fig1-1_scientific_method.jpg"
FIG2_IMG = ROOT / "Images" / "Chapter_01" / "embedded" / "fig1-2_accuracy_precision.png"

FIG1_KEY = "Figure 1.1 — The iterative cycle of the scientific method."
FIG2_KEY = "Figure 1.2 — Accuracy versus precision visualized."

FIG1_ATTRIB = (
    "Adapted from OpenStax Biology 2e, Figure 1.6 (Clark, Douglas, Choi). "
    "Access for free at openstax.org. Licensed CC BY 4.0."
)
FIG2_ATTRIB = (
    "Accuracy and Precision by Arbeck, via Wikimedia Commons. "
    "Licensed CC BY 4.0 (https://creativecommons.org/licenses/by/4.0/)."
)


def find_figure_table(doc, key_prefix):
    for t in doc.tables:
        try:
            first_text = t.rows[0].cells[0].text.strip()
        except IndexError:
            continue
        if first_text.startswith(key_prefix):
            return t
    return None


def insert_image_block_before_table(doc, table, image_path, attribution, width_in=4.5):
    body = doc.element.body
    tbl_el = table._element

    # Build the image paragraph by adding it to the doc body (so python-docx
    # registers the image part), then move its element to right before the
    # table. Same for the attribution paragraph.
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(image_path), width=Inches(width_in))

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(attribution)
    cap_run.italic = True
    cap_run.font.size = Pt(9)

    # Move the two newly-appended paragraphs to immediately before the table.
    img_el = img_p._element
    cap_el = cap_p._element
    img_el.getparent().remove(img_el)
    cap_el.getparent().remove(cap_el)
    tbl_el.addprevious(img_el)
    tbl_el.addprevious(cap_el)


def main():
    if not SRC.exists():
        sys.exit(f"Missing source: {SRC}")
    for p in (FIG1_IMG, FIG2_IMG):
        if not p.exists():
            sys.exit(f"Missing image: {p}")

    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, BACKUP)

    doc = Document(str(SRC))

    t1 = find_figure_table(doc, FIG1_KEY)
    t2 = find_figure_table(doc, FIG2_KEY)
    if t1 is None:
        sys.exit("Could not find Figure 1.1 description table")
    if t2 is None:
        sys.exit("Could not find Figure 1.2 description table")

    insert_image_block_before_table(doc, t1, FIG1_IMG, FIG1_ATTRIB)
    insert_image_block_before_table(doc, t2, FIG2_IMG, FIG2_ATTRIB)

    doc.save(str(SRC))
    print(f"OK — wrote {SRC}")
    print(f"Backup at {BACKUP}")


if __name__ == "__main__":
    main()
