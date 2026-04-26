"""Insert sourced OER images above the FIGURE DESCRIPTION tables in chapters 2-10.

The figure-description blocks live as Word tables (shaded boxes). This script
inserts a centered picture paragraph and a 9-pt italic attribution paragraph
immediately before each target table. The FIGURE DESCRIPTION block is left
in place to serve as accessible alt-text per project convention.

Where no license-compatible image exists for a figure, we leave the
FIGURE DESCRIPTION block alone (per the manifest's "If You Cannot Find a
Suitable Image" guidance).
"""
from pathlib import Path
import shutil
import sys

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
BACKUP_DIR = ROOT / ".firecrawl" / "backups"

CHAPTER_FILES = {
    "02": "Chapter_02_Unit_Systems_and_Dimensional_Analysis.docx",
    "03": "Chapter_03_Basic_Concepts_About_Matter.docx",
    "04": "Chapter_04_Atoms_Molecules_Subatomic_Particles.docx",
    "05": "Chapter_05_Electronic_Structure_and_Periodicity.docx",
    "06": "Chapter_06_Chemical_Bonds.docx",
    "10": "Chapter_10_States_of_Matter.docx",
}

OPENSTAX_CHEM = (
    "Adapted from OpenStax Chemistry 2e, {fig} (Flowers, Theopold, Langley, "
    "Robinson). Access for free at openstax.org. Licensed CC BY 4.0."
)

# (figure_header_prefix, image_filename, attribution)
PLAN = {
    "02": [
        ("Figure 2.2 — Aligned temperature scales.",
         "fig2-2_temperature_scales.png",
         OPENSTAX_CHEM.format(fig="Figure 1.28")),
    ],
    "03": [
        ("Figure 3.1 — Particle pictures of solid, liquid, and gas.",
         "fig3-1_solid_liquid_gas.png",
         OPENSTAX_CHEM.format(fig="Figure 1.6")),
        ("Figure 3.2 — Classification of matter (flowchart).",
         "fig3-2_matter_classification.png",
         OPENSTAX_CHEM.format(fig="Figure 1.11")),
    ],
    "04": [
        ("Figure 4.1 — Modern model of the atom (not to scale).",
         "fig4-1_modern_atom.png",
         OPENSTAX_CHEM.format(fig="Figure 2.11")),
        ("Figure 4.2 — Rutherford gold-foil experiment.",
         "fig4-2_rutherford_apparatus.png",
         OPENSTAX_CHEM.format(fig="Figure 2.9")),
    ],
    "05": [
        ("Figure 5.2 — Shapes of s, p, and d orbitals.",
         "fig5-2_orbital_shapes.png",
         OPENSTAX_CHEM.format(fig="Figure 6.21")),
        ("Figure 5.3 — Aufbau filling order (diagonal-arrow diagram).",
         "fig5-3_aufbau.png",
         OPENSTAX_CHEM.format(fig="Figure 6.26")),
        ("Figure 5.4 — Periodic trend in atomic radius.",
         "fig5-4_atomic_radius.png",
         OPENSTAX_CHEM.format(fig="Figure 6.31")),
    ],
    "06": [
        ("Figure 6.1 — VSEPR geometries for 2, 3, and 4 electron groups.",
         "fig6-1_vsepr.png",
         OPENSTAX_CHEM.format(fig="Figure 7.16")),
    ],
    "10": [
        ("Figure 10.2 — Hydrogen-bond network in liquid water vs. ice.",
         "fig10-2_hbond_water.png",
         OPENSTAX_CHEM.format(fig="Figure 10.10")),
        ("Figure 10.3 — Phase diagram of water (schematic, not to scale).",
         "fig10-3_phase_diagram_water.png",
         OPENSTAX_CHEM.format(fig="Figure 10.31")),
    ],
}


def find_figure_table(doc, key_prefix):
    for t in doc.tables:
        try:
            first_text = t.rows[0].cells[0].text.strip()
        except IndexError:
            continue
        if first_text.startswith(key_prefix):
            return t
    return None


def insert_image_block_before_table(doc, table, image_path, attribution, width_in=4.5):
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(image_path), width=Inches(width_in))

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(attribution)
    cap_run.italic = True
    cap_run.font.size = Pt(9)

    img_el = img_p._element
    cap_el = cap_p._element
    img_el.getparent().remove(img_el)
    cap_el.getparent().remove(cap_el)
    table._element.addprevious(img_el)
    table._element.addprevious(cap_el)


def process_chapter(num, filename, items):
    src = ROOT / filename
    if not src.exists():
        sys.exit(f"Missing source: {src}")
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{src.stem}.before-images.docx"
    shutil.copy2(src, backup)

    doc = Document(str(src))
    inserted = []
    for header, img_name, attribution in items:
        img_path = ROOT / "Images" / f"Chapter_{num}" / "embedded" / img_name
        if not img_path.exists():
            sys.exit(f"Missing image: {img_path}")
        t = find_figure_table(doc, header)
        if t is None:
            sys.exit(f"  Ch {num}: could not find table for {header!r}")
        insert_image_block_before_table(doc, t, img_path, attribution)
        inserted.append(header)

    doc.save(str(src))
    print(f"Ch {num}: wrote {len(inserted)} image(s) -> {src.name}")
    for h in inserted:
        print(f"   - {h}")


def main():
    for num, items in PLAN.items():
        filename = CHAPTER_FILES[num]
        process_chapter(num, filename, items)


if __name__ == "__main__":
    main()
