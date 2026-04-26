"""Rewrite-engine variant that targets Worked Example "Step N — …" paragraphs.

The Solutions engine in ``rewrite_engine.py`` walks "Solution:" paragraphs in
document order. Worked Examples are different: each example has its own
"Step 3 — …" line (or wherever the math chain lives), and we want to
modify those *in place* rather than scan by index. So this engine takes a
list of (anchor_text, new_text_with_strikes, explanation) tuples and finds
each anchor by string match within the first 80 chars of a paragraph.

The replacement preserves the paragraph's existing pPr (style /
indentation) and replaces only the runs.
"""
from copy import deepcopy
from pathlib import Path
import re
import shutil
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
BACKUP_DIR = ROOT / ".firecrawl" / "backups"

W_PPR = qn("w:pPr")
STRIKE_RE = re.compile(r"~~(.+?)~~")


def _add_run(p, text, *, italic=False, struck=False):
    run = p.add_run(text)
    if italic:
        run.italic = True
    if struck:
        rPr = run._element.get_or_add_rPr()
        s_el = OxmlElement("w:strike")
        s_el.set(qn("w:val"), "true")
        rPr.append(s_el)
    return run


def _add_runs_from_markup(p, markup, *, italic=False):
    pos = 0
    for m in STRIKE_RE.finditer(markup):
        if m.start() > pos:
            _add_run(p, markup[pos:m.start()], italic=italic)
        _add_run(p, m.group(1), italic=italic, struck=True)
        pos = m.end()
    if pos < len(markup):
        _add_run(p, markup[pos:], italic=italic)


def _clear_runs(p):
    p_el = p._element
    for child in list(p_el):
        if child.tag != W_PPR:
            p_el.remove(child)


def _insert_paragraph_after(anchor):
    new_p = OxmlElement("w:p")
    src_pPr = anchor._element.find(W_PPR)
    if src_pPr is not None:
        new_p.append(deepcopy(src_pPr))
    anchor._element.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def apply_example_rewrites(filename, rewrites):
    """Rewrites is a list of (anchor_substring, new_step_markup, explanation).

    For each entry we find the first paragraph whose text starts with
    `anchor_substring`. We replace its runs with `new_step_markup` (which
    may use ~~unit~~ for strikethrough), and insert an italic
    explanation paragraph immediately after.

    Returns count of rewrites applied.
    """
    src = ROOT / filename
    if not src.exists():
        sys.exit(f"Missing: {src}")

    backup = BACKUP_DIR / f"{src.stem}.before-examples-rewrite.docx"
    if not backup.exists():
        BACKUP_DIR.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, backup)

    doc = Document(str(src))

    def all_paragraphs():
        """Iterate every paragraph in the document body AND inside tables.

        EXAMPLE blocks in the OER chapters are formatted as single-cell
        tables for visual framing, so doc.paragraphs (which only returns
        top-level body paragraphs) misses them. We yield top-level
        paragraphs first, then descend into every cell of every table
        (including nested tables, recursively)."""
        for p in doc.paragraphs:
            yield p
        # Recurse through tables.
        def walk_tables(tables):
            for t in tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p
                        yield from walk_tables(cell.tables)
        yield from walk_tables(doc.tables)

    applied = 0
    for anchor, new_markup, explanation in rewrites:
        target = None
        for p in all_paragraphs():
            text = (p.text or "").strip()
            if text.startswith(anchor):
                target = p
                break
        if target is None:
            print(f"WARN: anchor not found: {anchor[:60]}...")
            continue

        # Idempotency: skip if this paragraph already has strikethrough.
        if "w:strike" in target._element.xml:
            continue

        _clear_runs(target)
        _add_runs_from_markup(target, new_markup)

        explanation_p = _insert_paragraph_after(target)
        _add_runs_from_markup(explanation_p, explanation, italic=True)

        applied += 1

    doc.save(str(src))
    return applied
