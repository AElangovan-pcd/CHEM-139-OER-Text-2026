"""Generic rewrite engine for "Solution:" blocks in any chapter.

Reads a per-chapter REWRITES list (sequence of (anchor_text, consume_after,
[math_lines], explanation) tuples, identical in shape to the Chapter 2
pilot) and applies them in order to the matching .docx file. Math lines
use the ~~unit~~ markup to mean Word strikethrough. The italic explanation
paragraph follows the math lines.

This is the same machinery as ``rewrite_solutions_ch02.py`` but factored
so each chapter only needs to ship a small data file with its REWRITES.
"""
from copy import deepcopy
from pathlib import Path
import re
import shutil
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
BACKUP_DIR = ROOT / ".firecrawl" / "backups"

W_T = qn("w:t")
W_PPR = qn("w:pPr")
STRIKE_RE = re.compile(r"~~(.+?)~~")


def _add_run(p, text, *, bold=False, italic=False, struck=False):
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if struck:
        rPr = run._element.get_or_add_rPr()
        s_el = OxmlElement("w:strike")
        s_el.set(qn("w:val"), "true")
        rPr.append(s_el)
    return run


def _add_runs_from_markup(p, markup, *, italic=False):
    pos = 0
    for m in STRIKE_RE.finditer(markup):
        if m.start() > pos:
            _add_run(p, markup[pos:m.start()], italic=italic)
        _add_run(p, m.group(1), italic=italic, struck=True)
        pos = m.end()
    if pos < len(markup):
        _add_run(p, markup[pos:], italic=italic)


def _clear_runs(p):
    p_el = p._element
    for child in list(p_el):
        if child.tag != W_PPR:
            p_el.remove(child)


def _insert_paragraph_after(anchor):
    new_p = OxmlElement("w:p")
    src_pPr = anchor._element.find(W_PPR)
    if src_pPr is not None:
        new_p.append(deepcopy(src_pPr))
    anchor._element.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def _is_solution(p):
    return (p.text or "").strip().startswith("Solution:")


def apply_rewrites(filename, rewrites, *, skip_indices=None):
    """Apply REWRITES (in document order of "Solution:" paragraphs) to the
    given chapter docx file. ``skip_indices`` is a set of 0-based
    Solution-paragraph indices to skip (used for purely qualitative answers
    that have no math to rewrite).

    Backs up to .firecrawl/backups/<stem>.before-solution-rewrite.docx.
    Returns the count of rewrites applied.
    """
    skip_indices = skip_indices or set()
    src = ROOT / filename
    if not src.exists():
        print(f"Missing: {src}")
        return 0

    backup = BACKUP_DIR / f"{src.stem}.before-solution-rewrite.docx"
    if not backup.exists():
        BACKUP_DIR.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, backup)

    doc = Document(str(src))
    paragraphs = list(doc.paragraphs)
    sol_index = 0
    rewrite_index = 0
    i = 0
    applied = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        if not _is_solution(p):
            i += 1
            continue

        if sol_index in skip_indices:
            sol_index += 1
            i += 1
            continue

        if rewrite_index >= len(rewrites):
            print(
                f"WARN: encountered Solution #{sol_index} in {filename} but only "
                f"{len(rewrites)} rewrites defined; stopping."
            )
            break

        anchor_text, consume_after, lines, explanation = rewrites[rewrite_index]

        body_el = p._element.getparent()
        nxt = p._element.getnext()
        consumed = 0
        while consumed < consume_after and nxt is not None:
            sibling = nxt
            nxt = nxt.getnext()
            body_el.remove(sibling)
            consumed += 1

        _clear_runs(p)
        _add_run(p, "Solution: ", bold=True)
        _add_runs_from_markup(p, lines[0])

        anchor = p
        for extra in lines[1:]:
            new_p = _insert_paragraph_after(anchor)
            _add_runs_from_markup(new_p, extra)
            anchor = new_p

        new_p = _insert_paragraph_after(anchor)
        _add_runs_from_markup(new_p, explanation, italic=True)

        paragraphs = list(doc.paragraphs)
        i = i + len(lines) + 1
        sol_index += 1
        rewrite_index += 1
        applied += 1

    doc.save(str(src))
    return applied
