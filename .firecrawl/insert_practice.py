"""Append an "Additional Practice Problems by Topic" section to each chapter
just before the "Multi-Concept Problems" heading, using accessible markup:

  * Section title is a real Heading 2 paragraph (semantic outline level).
  * Each topic title is a real Heading 3 paragraph (semantic outline level).
  * Each problem is a *real* numbered-list item: a paragraph carrying
    <w:numPr> referencing a decimal-format abstract numbering definition.
    Screen readers announce these as a list with item numbers, satisfying
    WCAG 2.2 Success Criterion 1.3.1 (Info and Relationships) — the
    numbering is encoded structurally, not just typed as "1. ".
  * Each topic gets its OWN <w:num> instance so numbering restarts at 1
    inside each topic.
  * Each solution is a continuation paragraph immediately after its problem.
    The label "Solution:" is set in BOLD (in addition to italic) so that
    emphasis is conveyed by more than a single visual style — never relying
    on color alone.
  * Heading hierarchy: H2 -> H3 -> body -> H3 -> body ... no skipped levels.

Idempotent: re-running detects the existing section and skips the chapter.
"""
from copy import deepcopy
from pathlib import Path
import shutil
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent))
from additional_problems import PROBLEMS

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
BACKUP_DIR = ROOT / ".firecrawl" / "backups"

CHAPTER_FILES = {
    "01": "Chapter_01_Science_and_Measurement.docx",
    "02": "Chapter_02_Unit_Systems_and_Dimensional_Analysis.docx",
    "03": "Chapter_03_Basic_Concepts_About_Matter.docx",
    "04": "Chapter_04_Atoms_Molecules_Subatomic_Particles.docx",
    "05": "Chapter_05_Electronic_Structure_and_Periodicity.docx",
    "06": "Chapter_06_Chemical_Bonds.docx",
    "07": "Chapter_07_Chemical_Nomenclature.docx",
    "08": "Chapter_08_Mole_Concept_and_Chemical_Formulas.docx",
    "09": "Chapter_09_Chemical_Calculations_and_Equations.docx",
    "10": "Chapter_10_States_of_Matter.docx",
}

NEW_SECTION_TITLE = "Additional Practice Problems by Topic"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------- numbering helpers (build accessible ordered lists) ------------ #


def get_numbering_xml(doc):
    """Return the docx numbering part's XML element."""
    return doc.part.numbering_part.element


def find_decimal_abstract_num_id(numbering_el):
    """Return an existing abstractNumId whose level-0 format is decimal.
    Falls back to None if not found."""
    for an in numbering_el.findall(qn("w:abstractNum")):
        lvl = an.find(qn("w:lvl"))
        if lvl is None:
            continue
        fmt = lvl.find(qn("w:numFmt"))
        if fmt is not None and fmt.get(qn("w:val")) == "decimal":
            return an.get(qn("w:abstractNumId"))
    return None


def add_decimal_abstract_num(numbering_el):
    """Append a fresh decimal abstractNum and return its id (string)."""
    used_ids = {int(a.get(qn("w:abstractNumId")))
                for a in numbering_el.findall(qn("w:abstractNum"))}
    new_id = max(used_ids, default=-1) + 1

    an = OxmlElement("w:abstractNum")
    an.set(qn("w:abstractNumId"), str(new_id))

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    txt = OxmlElement("w:lvlText")
    txt.set(qn("w:val"), "%1.")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")

    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)

    for child in (start, fmt, txt, jc, pPr):
        lvl.append(child)
    an.append(lvl)

    # Insert before the first <w:num> (or at end if none).
    first_num = numbering_el.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(an)
    else:
        numbering_el.append(an)
    return str(new_id)


def add_num_instance(numbering_el, abstract_num_id):
    """Add a fresh <w:num> referencing the given abstractNumId.
    Returns the new numId (string). Each topic gets its own so numbering
    restarts at 1 within each topic block."""
    used_ids = {int(n.get(qn("w:numId")))
                for n in numbering_el.findall(qn("w:num"))}
    new_id = max(used_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    aref = OxmlElement("w:abstractNumId")
    aref.set(qn("w:val"), str(abstract_num_id))
    num.append(aref)

    # Force a startOverride so this numId restarts at 1 even if abstractNum
    # was used earlier in the document.
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)

    numbering_el.append(num)
    return str(new_id)


def apply_numPr(paragraph, num_id, ilvl=0):
    """Mark a paragraph as a numbered-list item."""
    pPr = paragraph._p.get_or_add_pPr()
    # Remove any existing numPr.
    for existing in pPr.findall(qn("w:numPr")):
        pPr.remove(existing)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_el = OxmlElement("w:numId")
    num_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(num_el)
    pPr.append(numPr)


# ---------- insertion ---------------------------------------------------- #


def find_target_paragraph(doc, target_text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(target_text):
            return p
    return None


def already_present(doc):
    return find_target_paragraph(doc, NEW_SECTION_TITLE) is not None


def style_if_available(doc, name):
    """Resolve a paragraph-style by display name. Some chapter docx files
    contain duplicate style-name entries that defeat docx.styles[name]; in
    that case fall back to looking up by style_id (the no-space form,
    e.g. "Heading2"). Returns the resolved Style object, or None."""
    # 1. Direct lookup by display name.
    try:
        return doc.styles[name]
    except KeyError:
        pass
    # 2. Iterate and match by .name (handles ambiguous styles dict).
    for s in doc.styles:
        try:
            if s.name == name:
                return s
        except Exception:
            continue
    # 3. Try the no-space style_id form.
    try:
        return doc.styles[name.replace(" ", "")]
    except KeyError:
        return None


def insert_section_before(doc, anchor_p, chapter_num):
    h2_style = style_if_available(doc, "Heading 2")
    h3_style = style_if_available(doc, "Heading 3")
    list_para_style = style_if_available(doc, "List Paragraph")

    numbering_el = get_numbering_xml(doc)
    abstract_id = find_decimal_abstract_num_id(numbering_el)
    if abstract_id is None:
        abstract_id = add_decimal_abstract_num(numbering_el)

    def new_para(text="", style=None):
        # Insert without style first, then assign by style object if given.
        # This avoids the python-docx string-resolution path that fails on
        # docs with duplicate style-name entries.
        np = anchor_p.insert_paragraph_before(text)
        if style is not None:
            np.style = style
        return np

    # Section heading (H2)
    new_para(NEW_SECTION_TITLE, style=h2_style)
    new_para(
        "Five additional problems per topic, each followed by a complete "
        "worked solution. Cover the solution and try the problem first; then check."
    )

    topics = PROBLEMS.get(chapter_num, {})
    if not topics:
        return False

    for topic, items in topics.items():
        # H3 for each topic — preserves a clean H2 -> H3 outline.
        new_para(topic, style=h3_style)

        # One numId per topic so numbering restarts at 1 inside each block.
        topic_num_id = add_num_instance(numbering_el, abstract_id)

        for q, s in items:
            # Problem paragraph: real ordered-list item (no manual "1. ").
            qp = new_para(q, style=list_para_style)
            apply_numPr(qp, topic_num_id, ilvl=0)

            # Solution paragraph: continuation; "Solution:" in bold AND italic
            # so emphasis is not conveyed by a single visual style alone.
            sp = anchor_p.insert_paragraph_before("")
            run_label = sp.add_run("Solution: ")
            run_label.bold = True
            run_label.italic = True
            sp.add_run(s)

    return True


def process_chapter(num, filename):
    src = ROOT / filename
    if not src.exists():
        sys.exit(f"Missing: {src}")
    doc = Document(str(src))

    if already_present(doc):
        print(f"Ch {num}: section already present; skipping.")
        return

    anchor = find_target_paragraph(doc, "Multi-Concept Problems")
    if anchor is None:
        for alt in ("Multiple-Choice Practice Test", "Solutions to Practice Problems"):
            anchor = find_target_paragraph(doc, alt)
            if anchor is not None:
                break
    if anchor is None:
        sys.exit(f"Ch {num}: no anchor paragraph found.")

    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"{src.stem}.before-practice-a11y.docx"
    shutil.copy2(src, backup)

    inserted = insert_section_before(doc, anchor, num)
    if inserted:
        doc.save(str(src))
        topic_count = len(PROBLEMS.get(num, {}))
        print(f"Ch {num}: appended {topic_count} topic(s) x 5 problems.")
    else:
        print(f"Ch {num}: no problems data; nothing done.")


def main():
    for num, filename in CHAPTER_FILES.items():
        process_chapter(num, filename)


if __name__ == "__main__":
    main()
