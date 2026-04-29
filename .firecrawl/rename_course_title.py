"""Rename the course title across the three .docx files that contain it.

  "Introduction to Chemical Principles"  ->  "General Chemistry Prep"

Affected files (verified by prior grep):
  Front_Matter_Preface_License.docx                       (2 occurrences)
  Chapter_01_Science_and_Measurement.docx                 (1 occurrence)
  Chapter_02_Unit_Systems_and_Dimensional_Analysis.docx   (1 occurrence)

The replacement is done at the python-docx run level. We walk every paragraph
in the body and every paragraph inside every table cell (recursively), and
replace runs whose .text contains the old title. The earlier audit confirmed
the title appears as a contiguous string in word/document.xml in all three
files - i.e. it lives in a single Word run - so a run-level replace works
without merging adjacent runs.

Originals are copied to .firecrawl/backups/ before any edit.
"""
from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.table import Table

ROOT = Path(__file__).resolve().parent.parent
BACKUPS = ROOT / ".firecrawl" / "backups"

OLD = "Introduction to Chemical Principles"
NEW = "General Chemistry Prep"

TARGETS = [
    "Front_Matter_Preface_License.docx",
    "Chapter_01_Science_and_Measurement.docx",
    "Chapter_02_Unit_Systems_and_Dimensional_Analysis.docx",
]


def iter_paragraphs(doc):
    """Yield every paragraph in the document, including those nested inside
    table cells (and tables nested inside cells)."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for t in cell.tables:
                yield from _iter_table_paragraphs(t)


def replace_in_doc(doc) -> int:
    """Replace OLD with NEW in every run; return number of runs modified."""
    n = 0
    for p in iter_paragraphs(doc):
        for run in p.runs:
            if OLD in run.text:
                run.text = run.text.replace(OLD, NEW)
                n += 1
    return n


def main() -> None:
    BACKUPS.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    total = 0
    for name in TARGETS:
        src = ROOT / name
        if not src.exists():
            print(f"SKIP missing: {name}")
            continue
        backup = BACKUPS / f"{src.stem}.{stamp}.bak.docx"
        shutil.copy2(src, backup)

        doc = Document(str(src))
        n = replace_in_doc(doc)
        if n == 0:
            print(f"  {name}: no matching runs found - leaving unchanged")
            continue
        doc.save(str(src))
        total += n
        print(f"  {name}: replaced in {n} run(s); backup -> {backup.name}")

    print(f"\nDone. {total} run replacement(s) total across {len(TARGETS)} files.")


if __name__ == "__main__":
    main()
