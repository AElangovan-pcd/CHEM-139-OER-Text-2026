"""Insert the original Fig 5.1 periodic-table PNG above its FIGURE DESCRIPTION
table in Chapter 5. Includes a duplicate-detection guard so re-running is safe.
"""
from pathlib import Path
import shutil

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
SRC = ROOT / "Chapter_05_Electronic_Structure_and_Periodicity.docx"
BACKUP = ROOT / ".firecrawl" / "backups" / f"{SRC.stem}.before-fig5-1.docx"
IMG = ROOT / "Images" / "Chapter_05" / "embedded" / "fig5-1_periodic_table.png"
KEY = "Figure 5.1 — The periodic table annotated with"
ATTRIB = (
    "Original figure for the CHEM 139 OER textbook (2026 ed.). "
    "Built from public IUPAC 2021 / NIST element data. Licensed CC BY 4.0."
)
NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def find_table(doc, key):
    for t in doc.tables:
        try:
            txt = t.rows[0].cells[0].text.strip()
        except IndexError:
            continue
        if txt.startswith(key):
            return t
    return None


def already_inserted(table):
    prev = table._element.getprevious()
    for _ in range(3):
        if prev is None:
            return False
        if prev.tag.endswith('}p') and prev.find(f'.//{NS}drawing') is not None:
            return True
        prev = prev.getprevious()
    return False


def main():
    BACKUP.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, BACKUP)

    doc = Document(str(SRC))
    t = find_table(doc, KEY)
    if t is None:
        raise SystemExit("Could not find Figure 5.1 description table")
    if already_inserted(t):
        print("Already inserted; skipping.")
        return

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(IMG), width=Inches(6.0))

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(ATTRIB)
    cap_run.italic = True
    cap_run.font.size = Pt(9)

    img_el = img_p._element
    cap_el = cap_p._element
    img_el.getparent().remove(img_el)
    cap_el.getparent().remove(cap_el)
    t._element.addprevious(img_el)
    t._element.addprevious(cap_el)

    doc.save(str(SRC))
    print(f"Inserted Fig 5.1 into {SRC.name}")
    print(f"Backup at {BACKUP}")


if __name__ == "__main__":
    main()
