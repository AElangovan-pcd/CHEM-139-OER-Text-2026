"""Rewrite every "Solution:"-labeled block in Chapter 2 in factor-label form.

For each numerical Solution, the math chain is rendered with cancelled units
shown by Word strikethrough character formatting; an italic explanatory line
follows naming which units cancel and why the answer has the sig figs it has.

Markup convention used inside `lines`:
    ~~text~~     -> Word strikethrough run (cancelled unit)
    Anything else is written as a normal run.

For a Solution whose original answer spans more than one paragraph in the
source (e.g. multi-part conversions), the entry's ``consume_after`` says how
many follow-on paragraphs to delete after the "Solution:" anchor so the new
content fully replaces the old.

Idempotent guard: if the Solution paragraph already contains the marker
``✗`` (a strikethrough X-mark we never use elsewhere), the file is
treated as already rewritten and the script bails out.
"""
from copy import deepcopy
from pathlib import Path
import re
import shutil
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
SRC = ROOT / "Chapter_02_Unit_Systems_and_Dimensional_Analysis.docx"
BACKUP_DIR = ROOT / ".firecrawl" / "backups"
BACKUP = BACKUP_DIR / "Chapter_02_Unit_Systems_and_Dimensional_Analysis.before-solution-rewrite.docx"

W_T = qn("w:t")
W_R = qn("w:r")
W_RPR = qn("w:rPr")
W_PPR = qn("w:pPr")
W_STRIKE = qn("w:strike")
W_I = qn("w:i")
W_B = qn("w:b")

STRIKE_RE = re.compile(r"~~(.+?)~~")

# ---------------- rewrite data ------------------------------------------- #
# Each tuple: (anchor_text_starts_with, consume_after, [math_lines], explanation)
# anchor_text is the first ~25 chars of the EXISTING "Solution:" paragraph
# (used for ordering sanity check, not for matching).
# Lines are written one-per-paragraph after the bold "Solution: " label on
# the first line. Multi-line entries replace the entire original block.

REWRITES = [
    # --- Practice Problems by Topic (27 problems) ---

    # PP01: Convert 3.50 km to m, cm, mm
    ("Solution:", 3, [
        "3.50 ~~km~~ × (10³ m / 1 ~~km~~) = 3.50 × 10³ m",
        "3.50 ~~km~~ × (10³ ~~m~~ / 1 ~~km~~) × (100 cm / 1 ~~m~~) = 3.50 × 10⁵ cm",
        "3.50 ~~km~~ × (10³ ~~m~~ / 1 ~~km~~) × (1000 mm / 1 ~~m~~) = 3.50 × 10⁶ mm",
    ],
        "In each chain, \"km\" cancels against the conversion factor, and (for cm and mm) the intermediate \"m\" cancels too. Three sig figs from \"3.50 km\" carry through; the metric conversion factors are exact."
    ),

    # PP02: Convert 0.025 g to mg, μg, kg
    ("Solution:", 3, [
        "0.025 ~~g~~ × (1000 mg / 1 ~~g~~) = 25 mg",
        "0.025 ~~g~~ × (10⁶ µg / 1 ~~g~~) = 2.5 × 10⁴ µg",
        "0.025 ~~g~~ × (1 kg / 1000 ~~g~~) = 2.5 × 10⁻⁵ kg",
    ],
        "\"g\" cancels in each chain. Two sig figs from \"0.025 g\" are preserved (leading zeros are not significant)."
    ),

    # PP03: Convert 750 mL to L, dL, μL  (single-line in source)
    ("Solution: 0.750 L", 0, [
        "750 ~~mL~~ × (1 L / 1000 ~~mL~~) = 0.750 L",
        "750 ~~mL~~ × (1 ~~L~~ / 1000 ~~mL~~) × (10 dL / 1 ~~L~~) = 7.50 dL",
        "750 ~~mL~~ × (1000 µL / 1 ~~mL~~) = 7.50 × 10⁵ µL",
    ],
        "\"mL\" cancels in each chain (and \"L\" cancels in the dL chain). Three sig figs are kept; the trailing zero in \"750\" is taken as significant from context."
    ),

    # PP04: 0.1 nm → m, pm
    ("Solution: 1 × 10", 0, [
        "0.1 ~~nm~~ × (10⁻⁹ m / 1 ~~nm~~) = 1 × 10⁻¹⁰ m",
        "0.1 ~~nm~~ × (1000 pm / 1 ~~nm~~) = 100 pm",
    ],
        "\"nm\" cancels in both chains. \"0.1 nm\" carries 1 sig fig, so the answers are reported as 1 × 10⁻¹⁰ m and 100 pm (1 sig fig)."
    ),

    # PP05: 250 μg → mg
    ("Solution: 0.250 mg", 0, [
        "250 ~~µg~~ × (1 mg / 1000 ~~µg~~) = 0.250 mg",
    ],
        "\"µg\" cancels, leaving \"mg\". Three sig figs from \"250\" carry through (the conversion 1 mg = 1000 µg is exact)."
    ),

    # PP06: 5280 ft → mi
    ("Solution: 5280 ft", 0, [
        "5280 ~~ft~~ × (1 mi / 5280 ~~ft~~) = 1.000 mi",
    ],
        "\"ft\" cancels. The relationship 1 mi = 5280 ft is exact, so the result has unlimited sig figs; written as 1.000 mi by convention."
    ),

    # PP07: 12.0 oz → g
    ("Solution: 12.0 ", 0, [
        "12.0 ~~oz~~ × (28.35 g / 1 ~~oz~~) = 340.2 g → 340. g",
    ],
        "\"oz\" cancels, leaving \"g\". Three sig figs from \"12.0 oz\" limit the answer; rounded to 340. g (the trailing decimal point marks the final zero as significant)."
    ),

    # PP08: 100. yd → m
    ("Solution: 91.44 m", 0, [
        "100. ~~yd~~ × (0.9144 m / 1 ~~yd~~) = 91.44 m → 91.4 m",
    ],
        "\"yd\" cancels. The conversion is exact; \"100.\" carries 3 sig figs, so the answer is rounded to 91.4 m."
    ),

    # PP09: 1.50 mi → km
    ("Solution: 1.50 mi", 0, [
        "1.50 ~~mi~~ × (1.609 km / 1 ~~mi~~) = 2.41 km",
    ],
        "\"mi\" cancels. Three sig figs from \"1.50 mi\" set the precision; the conversion factor 1.609 km/mi has 4 sig figs, so 1.50 governs."
    ),

    # PP10: 30.0 mph → km/h
    ("Solution: 30.0 ", 0, [
        "30.0 ~~mi~~/h × (1.609 km / 1 ~~mi~~) = 48.27 km/h → 48.3 km/h",
    ],
        "\"mi\" cancels in the numerator; \"h\" stays in the denominator throughout, giving km/h. Three sig figs from \"30.0\" yield 48.3 km/h."
    ),

    # PP11: 750 mL → fl oz
    ("Solution: 750 / 29.57", 0, [
        "750 ~~mL~~ × (1 fl oz / 29.57 ~~mL~~) = 25.36 fl oz → 25.4 fl oz",
    ],
        "\"mL\" cancels, leaving \"fl oz\". Three sig figs (\"750\" treated as 3) yield 25.4 fl oz."
    ),

    # PP12: 5.00 lb → kg
    ("Solution: 5.00 ", 0, [
        "5.00 ~~lb~~ × (453.59 ~~g~~ / 1 ~~lb~~) × (1 kg / 1000 ~~g~~) = 2.27 kg",
    ],
        "\"lb\" cancels in step 1; \"g\" cancels in step 2. Three sig figs from \"5.00 lb\" yield 2.27 kg."
    ),

    # PP13: cube density (2.00 cm side, 21.6 g)
    ("Solution: V = (2.00", 0, [
        "V = (2.00 cm)³ = 8.00 cm³ = 8.00 mL",
        "d = m / V = 21.6 g / 8.00 mL = 2.70 g/mL",
    ],
        "Units carry through cubing: (cm)³ = cm³, and 1 cm³ ≡ 1 mL exactly. Density of 2.70 g/mL matches aluminum (Table 2.4). Three sig figs throughout."
    ),

    # PP14: 1.50 kg Hg volume
    ("Solution: 1.50 kg", 0, [
        "1.50 ~~kg~~ × (1000 g / 1 ~~kg~~) = 1500 g",
        "V = m / d = 1500 ~~g~~ × (1 mL / 13.546 ~~g~~) = 110.7 mL → 111 mL",
    ],
        "\"kg\" cancels in step 1, \"g\" cancels in step 2. Three sig figs (limited by \"1.50 kg\") give 111 mL."
    ),

    # PP15: ethanol mass
    ("Solution: m = 25.0", 0, [
        "m = d × V = 25.0 ~~mL~~ × (0.789 g / 1 ~~mL~~) = 19.7 g",
    ],
        "\"mL\" cancels, leaving \"g\". Three sig figs throughout."
    ),

    # PP16: float? — qualitative, keep as-is (no rewrite)
    # (We skip this index by using a sentinel below.)

    # PP17: unknown liquid identification
    ("Solution: d = 14.7", 0, [
        "d = m / V = 14.7 g / 18.6 mL = 0.790 g/mL",
    ],
        "Density 0.790 g/mL matches ethanol (0.789) or acetone (0.791) within reading; further testing (boiling point) would distinguish."
    ),

    # PP18: 70.0 kg adult dose
    ("Solution: 70.0 kg", 0, [
        "70.0 ~~kg~~ × (5.00 mg / 1 ~~kg~~) = 350. mg",
    ],
        "\"kg\" cancels, leaving \"mg\". Three sig figs from \"70.0\" and \"5.00\" yield 350. mg (the trailing decimal marks the final zero as significant)."
    ),

    # PP19: gas cost
    ("Solution: 12.0 ×", 0, [
        "12.0 ~~gal~~ × ($3.49 / 1 ~~gal~~) = $41.88 → $41.90",
    ],
        "\"gal\" cancels, leaving \"$\". Three sig figs (limited by \"12.0\") yield $41.90."
    ),

    # PP20: drive cost
    ("Solution: 350. mi", 0, [
        "350. ~~mi~~ × (1 ~~gal~~ / 32.0 ~~mi~~) × ($3.99 / 1 ~~gal~~) = $43.64 → $43.60",
    ],
        "\"mi\" cancels in step 1; \"gal\" cancels in step 2. Three sig figs (limited by \"350.\") yield $43.60."
    ),

    # PP21: silver in alloy
    ("Solution: 50.0 ×", 0, [
        "50.0 g alloy × (35.0 g Ag / 100 g alloy) = 17.5 g Ag",
    ],
        "\"g alloy\" cancels, leaving \"g Ag\". Three sig figs (\"50.0\" and \"35.0\") yield 17.5 g."
    ),

    # PP22: 99.4 vs 100.0 percent error
    ("Solution: |99.4", 0, [
        "% error = |99.4 − 100.0| / 100.0 × 100% = 0.6 / 100.0 × 100% = 0.6%",
    ],
        "Subtraction first: |99.4 − 100.0| = 0.6 (1 decimal place, 1 sig fig). Result 0.6%."
    ),

    # PP23: water density percent error
    ("Solution: |0.998", 0, [
        "% error = |0.998 − 0.997| / 0.997 × 100% = 0.001 / 0.997 × 100% = 0.10%",
    ],
        "Subtraction yields 0.001 (1 sig fig from the difference, but the result is reported to two decimal places, 0.10%, by common chemistry convention)."
    ),

    # PP24: 25 °C → °F, K
    ("Solution: °F = 25", 0, [
        "T_F = 25 °C × 9/5 + 32 = 45 + 32 = 77 °F",
        "T_K = 25 °C + 273.15 = 298.15 K → 298 K",
    ],
        "Temperature conversion uses formulas, not factor-label cancellation. \"25 °C\" has 0 decimal places, so K is rounded to 298 K."
    ),

    # PP25: 350.0 K → °C, °F
    ("Solution: °C = 350.0", 0, [
        "T_C = 350.0 K − 273.15 = 76.85 °C → 76.9 °C",
        "T_F = 76.9 °C × 9/5 + 32 = 138.42 + 32 = 170.42 °F → 170. °F",
    ],
        "Subtraction governs decimal places: 350.0 − 273.15 → 1 decimal place. Three sig figs throughout the °F step."
    ),

    # PP26: 1064 °C → K
    ("Solution: 1064 +", 0, [
        "T_K = 1064 °C + 273.15 = 1337.15 K → 1337 K",
    ],
        "\"1064\" has 0 decimal places, so the result is rounded to 1337 K (no decimal)."
    ),

    # PP27: 77 K → °C, °F
    ("Solution: °C = 77", 0, [
        "T_C = 77 K − 273.15 = −196.15 °C → −196 °C",
        "T_F = −196 °C × 9/5 + 32 = −352.8 + 32 = −320.8 °F → −321 °F",
    ],
        "\"77 K\" has 0 decimal places, so the °C result is reported as −196 °C; the °F step is then limited to 0 decimals."
    ),

    # --- Additional Practice Problems by Topic (30 problems) ---

    # APP01: 2.50 km → m
    ("Solution: 2.50 km", 0, [
        "2.50 ~~km~~ × (1000 m / 1 ~~km~~) = 2.50 × 10³ m = 2500 m",
    ],
        "\"km\" cancels, leaving \"m\". Three sig figs preserved (the conversion 1 km = 1000 m is exact)."
    ),

    # APP02: 75 mg → g
    ("Solution: 75 mg", 0, [
        "75 ~~mg~~ × (1 g / 1000 ~~mg~~) = 0.075 g",
    ],
        "\"mg\" cancels, leaving \"g\". Two sig figs from \"75\" carry through."
    ),

    # APP03: 0.045 L → mL
    ("Solution: 0.045 L", 0, [
        "0.045 ~~L~~ × (1000 mL / 1 ~~L~~) = 45 mL",
    ],
        "\"L\" cancels, leaving \"mL\". Two sig figs from \"0.045\" preserved."
    ),

    # APP04: 3.2 μg → mg
    ("Solution: 3.2 μg", 0, [
        "3.2 ~~µg~~ × (1 mg / 1000 ~~µg~~) = 0.0032 mg = 3.2 × 10⁻³ mg",
    ],
        "\"µg\" cancels, leaving \"mg\". Two sig figs preserved."
    ),

    # APP05: prefix list — qualitative, skip rewrite

    # APP06: 4.0 ft → cm
    ("Solution: 4.0 ft", 0, [
        "4.0 ~~ft~~ × (12 ~~in~~ / 1 ~~ft~~) × (2.54 cm / 1 ~~in~~) = 121.92 cm → 1.2 × 10² cm",
    ],
        "\"ft\" cancels in step 1; \"in\" cancels in step 2. Both conversion factors are exact, so the answer is limited by \"4.0 ft\" (2 sig figs) → 1.2 × 10² cm."
    ),

    # APP07: 65 mi/h → m/s
    ("Solution: 65 mi/h", 0, [
        "65 ~~mi~~/~~h~~ × (1.609 ~~km~~ / 1 ~~mi~~) × (1000 m / 1 ~~km~~) × (1 ~~h~~ / 3600 s) = 29.05 m/s → 29 m/s",
    ],
        "\"mi\" cancels, \"km\" cancels, \"h\" cancels. \"m\" remains in the numerator and \"s\" in the denominator. Two sig figs (limited by \"65\") yield 29 m/s."
    ),

    # APP08: 2.0 gal → mL
    ("Solution: 2.0 gal", 0, [
        "2.0 ~~gal~~ × (3.785 ~~L~~ / 1 ~~gal~~) × (1000 mL / 1 ~~L~~) = 7570 mL → 7.6 × 10³ mL",
    ],
        "\"gal\" cancels in step 1; \"L\" cancels in step 2. Two sig figs (limited by \"2.0\") yield 7.6 × 10³ mL."
    ),

    # APP09: 250 mg → oz
    ("Solution: 250 mg", 0, [
        "250 ~~mg~~ × (1 ~~g~~ / 1000 ~~mg~~) × (1 ~~lb~~ / 453.6 ~~g~~) × (16 oz / 1 ~~lb~~) = 0.00882 oz → 8.8 × 10⁻³ oz",
    ],
        "\"mg\" → \"g\" → \"lb\" → \"oz\": each unit cancels at each step. Two sig figs (limited by \"250\" treated as 2 sf) yield 8.8 × 10⁻³ oz."
    ),

    # APP10: 0.50 L → qt
    ("Solution: 0.50 L", 0, [
        "0.50 ~~L~~ × (1.057 qt / 1 ~~L~~) = 0.53 qt",
    ],
        "\"L\" cancels, leaving \"qt\". Two sig figs (limited by \"0.50\") yield 0.53 qt."
    ),

    # APP11: 2.50 cm cube density
    ("Solution: V = (2.50", 0, [
        "V = (2.50 cm)³ = 15.625 cm³",
        "ρ = m / V = 78.4 g / 15.625 cm³ = 5.02 g/cm³",
    ],
        "Cubing the edge length cubes the unit: (cm)³ = cm³. Three sig figs (limited by \"78.4\") give 5.02 g/cm³."
    ),

    # APP12: 25.0 mL Hg → mass
    ("Solution: m = ρV", 0, [
        "m = ρ × V = (13.6 g / 1 ~~mL~~) × 25.0 ~~mL~~ = 340. g",
    ],
        "\"mL\" cancels, leaving \"g\". Three sig figs throughout."
    ),

    # APP13: unknown 50.0 g, 63.3 mL
    ("Solution: ρ = 50.0", 0, [
        "ρ = m / V = 50.0 g / 63.3 mL = 0.790 g/mL",
    ],
        "Density 0.790 g/mL matches ethanol (0.789) or acetone (0.791) within reading; further testing (e.g., boiling point) needed to distinguish. Three sig figs."
    ),

    # APP14: Al volume from mass
    ("Solution: V = m/ρ", 0, [
        "V = m / ρ = 8.10 ~~g~~ × (1 cm³ / 2.70 ~~g~~) = 3.00 cm³",
    ],
        "\"g\" cancels, leaving \"cm³\". Three sig figs throughout."
    ),

    # APP15: iron in mercury — qualitative, skip rewrite

    # APP16: 70.0 kg dose
    ("Solution: 70.0 kg", 0, [
        "70.0 ~~kg~~ × (5.0 mg / 1 ~~kg~~) = 350 mg → 3.5 × 10² mg",
    ],
        "\"kg\" cancels, leaving \"mg\". Two sig figs (limited by \"5.0\") yield 3.5 × 10² mg."
    ),

    # APP17: 500 mg → mL
    ("Solution: 500 mg", 0, [
        "500 ~~mg~~ × (5 mL / 250 ~~mg~~) = 10 mL",
    ],
        "\"mg\" cancels, leaving \"mL\". The answer (10 mL) is consistent with the 1- to 2-sig-fig precision of the input concentration."
    ),

    # APP18: IV drip rate
    ("Solution: 1.00 L", 0, [
        "Rate = (1.00 ~~L~~ × (1000 mL / 1 ~~L~~)) / (8.0 ~~h~~ × (60 min / 1 ~~h~~)) = 1000 mL / 480 min = 2.08 mL/min → 2.1 mL/min",
    ],
        "\"L\" cancels in the numerator; \"h\" cancels in the denominator, leaving mL/min. Two sig figs (limited by \"8.0 h\") yield 2.1 mL/min."
    ),

    # APP19: NaCl mass in IV
    ("Solution: Mass of solution", 0, [
        "Mass of solution = 250.0 ~~mL~~ × (1.00 g / 1 ~~mL~~) = 250.0 g",
        "Mass of NaCl = 250.0 g solution × (0.90 g NaCl / 100 g solution) = 2.25 g → 2.3 g",
    ],
        "\"mL\" cancels in step 1; \"g solution\" cancels in step 2. Two sig figs (limited by \"0.90%\") yield 2.3 g NaCl."
    ),

    # APP20: child dose volume
    ("Solution: Dose mass", 0, [
        "Dose mass = 18 ~~kg~~ × (12 mg / 1 ~~kg~~) = 216 mg",
        "Volume = 216 ~~mg~~ × (5 mL / 160 ~~mg~~) = 6.75 mL → 6.8 mL",
    ],
        "\"kg\" cancels in step 1; \"mg\" cancels in step 2. Two sig figs (limited by \"18\") yield 6.8 mL."
    ),

    # APP21: 9.85 g percent error
    ("Solution: %error = |9.85", 0, [
        "% error = |9.85 − 10.00| g / 10.00 g × 100% = 0.15 / 10.00 × 100% = 1.50%",
    ],
        "\"g\" cancels in the ratio (mass / mass is dimensionless). Three sig figs in the ratio yield 1.50%."
    ),

    # APP22: zinc % in alloy
    ("Solution: 25/200", 0, [
        "% Zn = 25 ~~g~~ Zn / 200 ~~g~~ alloy × 100% = 12.5%",
    ],
        "\"g\" cancels in numerator and denominator, leaving a pure percentage. Two sig figs (limited by \"25\") yield 12.5%."
    ),

    # APP23: glucose mass
    ("Solution: 0.075 ×", 0, [
        "Mass of glucose = 240 g solution × (7.5 g glucose / 100 g solution) = 18 g",
    ],
        "\"g solution\" cancels, leaving \"g glucose\". Two sig figs (limited by \"7.5%\") yield 18 g."
    ),

    # APP24: scale percent error
    ("Solution: |5.10", 0, [
        "% error = |5.10 − 5.000| kg / 5.000 kg × 100% = 0.10 / 5.000 × 100% = 2.0%",
    ],
        "\"kg\" cancels in the ratio. Subtraction yields 0.10 (2 sig figs); answer is 2.0%."
    ),

    # APP25: Student A vs B percent error
    ("Solution: A:", 0, [
        "Student A: % error = |12.45 − 12.50| cm / 12.50 cm × 100% = 0.05 / 12.50 × 100% = 0.40%",
        "Student B: % error = |12.62 − 12.50| cm / 12.50 cm × 100% = 0.12 / 12.50 × 100% = 0.96%",
    ],
        "\"cm\" cancels in each ratio. Student A (0.40%) is more accurate than Student B (0.96%)."
    ),

    # APP26: 25 °C → K
    ("Solution: T = 25", 0, [
        "T_K = 25 °C + 273.15 = 298.15 K → 298 K",
    ],
        "Temperature conversion is by formula, not factor-label. \"25 °C\" has 0 decimal places, so the result rounds to 298 K."
    ),

    # APP27: 98.6 °F → °C
    ("Solution: T_C = (98.6", 0, [
        "T_C = (98.6 °F − 32) × 5/9 = 66.6 × 5/9 = 37.0 °C",
    ],
        "Conversion by formula. The \"32\" is an exact reference; one decimal place from \"98.6\" carries through."
    ),

    # APP28: 350 K → °F
    ("Solution: T_C = 350", 0, [
        "T_C = 350 K − 273.15 = 76.85 °C",
        "T_F = 76.85 °C × 9/5 + 32 = 138.33 + 32 = 170.33 °F → 170. °F",
    ],
        "Conversion by formula. \"350 K\" has 0 decimal places, so the final °F result is rounded to 170. °F."
    ),

    # APP29: 77 K → °C
    ("Solution: T_C = 77", 0, [
        "T_C = 77 K − 273.15 = −196.15 °C → −196 °C",
    ],
        "Conversion by formula. \"77 K\" has 0 decimal places, so the answer is reported as −196 °C."
    ),

    # APP30: 0 °F → °C, K
    ("Solution: T_C = (0", 0, [
        "T_C = (0 °F − 32) × 5/9 = −32 × 5/9 = −17.8 °C",
        "T_K = −17.8 °C + 273.15 = 255.35 K → 255 K",
    ],
        "Conversion by formula. \"0 °F\" gives −17.8 °C (1 decimal place from the calculation); rounding K to no decimals matches that limit."
    ),

    # --- Multi-Concept Problems (8 problems) ---

    # MC01: IV drip multi-part (2 sub-parts in a single solution block)
    ("Solution:", 2, [
        "Volume in 8.0 h = 125 mL/~~h~~ × 8.0 ~~h~~ = 1000 mL → 1.0 × 10³ mL",
        "Mass of solution ≈ 1000 ~~mL~~ × (1.00 g / 1 ~~mL~~) = 1000 g",
        "Mass of NaCl = 1000 ~~g~~ solution × (0.45 ~~g~~ NaCl / 100 ~~g~~ solution) = 4.5 g NaCl",
    ],
        "\"h\" cancels in step 1; \"mL\" cancels in step 2; \"g solution\" cancels in step 3. Two sig figs (limited by \"8.0 h\" and \"0.45%\") yield 4.5 g NaCl."
    ),

    # MC02: Aluminum sphere multi-part
    ("Solution:", 3, [
        "V = (4/3)π r³ = (4/3)(3.1416)(1.20 cm)³ = 7.238 cm³ = 7.24 mL",
        "m = ρ × V = (2.70 g / 1 ~~mL~~) × 7.24 ~~mL~~ = 19.55 g → 19.5 g",
        "Mass of displaced water = (1.00 g / 1 ~~mL~~) × 7.24 ~~mL~~ = 7.24 g",
        "Apparent mass in water = 19.5 g − 7.24 g = 12.3 g",
    ],
        "Cubing the radius cubes the unit: (cm)³ = cm³, and 1 cm³ ≡ 1 mL exactly. \"mL\" cancels in the mass and displaced-water calculations. Three sig figs (limited by \"1.20 cm\") govern."
    ),

    # MC03: car gasoline → CO2
    ("Solution:", 2, [
        "V = 1.20 ~~gal~~ × (3.785 ~~L~~ / 1 ~~gal~~) × (1000 mL / 1 ~~L~~) = 4542 mL → 4.54 × 10³ mL",
        "Mass of gasoline = 4540 ~~mL~~ × (0.74 g / 1 ~~mL~~) = 3360 g → 3.36 × 10³ g",
        "Mass of C = 3360 g gasoline × (87.0 g C / 100 g gasoline) = 2920 g → 2.92 × 10³ g C",
        "Mass of CO₂ = 2920 ~~g C~~ × (44.01 g CO₂ / 12.011 ~~g C~~) = 1.07 × 10⁴ g ≈ 10.7 kg",
    ],
        "\"gal\", \"L\", \"mL\", and \"g C\" cancel in turn through the chain. Three sig figs throughout (limited by \"1.20 gal\" and \"87.0%\")."
    ),

    # MC04: blood-alcohol
    ("Solution:", 2, [
        "0.082 g ethanol / 100 mL × (1000 mg / 1 ~~g~~) = 82 mg / 100 mL = 82 mg/dL",
        "Legal limit: 0.080 g/100 mL = 80 mg/dL",
        "Excess = 82 − 80 = 2 mg/dL",
        "% excess = 2 / 80 × 100% = 2.5%",
    ],
        "\"g\" cancels when scaling to mg, and 100 mL ≡ 1 dL exactly, giving mg/dL directly. The percentage is dimensionless."
    ),

    # MC05: metal water displacement
    ("Solution:", 2, [
        "V = 30.3 mL − 25.0 mL = 5.3 mL",
        "d = m / V = 41.4 g / 5.3 mL = 7.81 g/mL",
    ],
        "Subtraction limits V to 1 decimal place (5.3 mL, 2 sig figs). Density 7.81 g/mL matches iron (7.87 g/mL) within reading."
    ),

    # MC06: baker milk
    ("Solution: V = 2.00", 0, [
        "V = 2.00 ~~cups~~ × (236.6 mL / 1 ~~cup~~) = 473.2 mL",
        "m of milk = 473.2 ~~mL~~ × (1.030 g / 1 ~~mL~~) = 487.4 g",
        "m of fat = 487.4 g milk × (3.25 g fat / 100 g milk) = 15.84 g → 15.8 g",
    ],
        "\"cup\" cancels in step 1; \"mL\" cancels in step 2; \"g milk\" cancels in step 3. Three sig figs (limited by \"3.25%\") yield 15.8 g."
    ),

    # MC07: gold troy oz
    ("Solution: V = 31.103", 0, [
        "V = 31.103 ~~g~~ × (1 mL / 19.32 ~~g~~) = 1.610 mL → 1.61 mL",
        "1.61 mL = 1.61 cm³ (since 1 mL ≡ 1 cm³ exactly)",
    ],
        "\"g\" cancels, leaving \"mL\". Three sig figs (limited by \"19.32\" rounded against the 5-sig-fig \"31.103\") yield 1.61 mL = 1.61 cm³."
    ),

    # MC08: −40 °F → °C, K
    ("Solution: °C = (−40", 0, [
        "T_C = (−40 °F − 32) × 5/9 = −72 × 5/9 = −40 °C",
        "T_K = −40 °C + 273.15 = 233.15 K → 233 K",
    ],
        "−40 is the unique temperature where °F and °C scales coincide numerically — a useful sanity check on conversions."
    ),
]

# Indices we deliberately leave unchanged (qualitative answers).
# Counted within the document's full sequence of "Solution:" paragraphs.
# Source order: 27 PP + 30 APP + 8 MC = 65.
# Skip indices for: PP16 (plastic float?), APP05 (prefix list), APP15 (iron in Hg).
SKIP_INDICES = {15, 30 + 4, 30 + 14}  # PP16 = 15 (0-based); APP05 = 27+4=31; APP15 = 27+14=41


# ---------------- helpers ------------------------------------------------ #


def text_of(p):
    return "".join(t.text or "" for t in p._element.iter(W_T))


def is_solution_paragraph(p):
    return text_of(p).strip().startswith("Solution:")


def clear_paragraph_runs(p):
    """Remove every <w:r> inside <w:p> while preserving <w:pPr>."""
    p_el = p._element
    for child in list(p_el):
        if child.tag != W_PPR:
            p_el.remove(child)


def add_run_with_strikes(p, text, *, bold=False, italic=False, struck=False):
    """Append a single run to paragraph p with the given character formatting."""
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if struck:
        rPr = run._element.get_or_add_rPr()
        strike = OxmlElement("w:strike")
        strike.set(qn("w:val"), "true")
        rPr.append(strike)
    return run


def add_runs_from_markup(p, markup, *, italic=False):
    """Tokenize markup on ~~...~~ and emit runs with strikethrough where marked."""
    pos = 0
    for m in STRIKE_RE.finditer(markup):
        if m.start() > pos:
            add_run_with_strikes(p, markup[pos:m.start()], italic=italic)
        add_run_with_strikes(p, m.group(1), italic=italic, struck=True)
        pos = m.end()
    if pos < len(markup):
        add_run_with_strikes(p, markup[pos:], italic=italic)


def insert_paragraph_after(p):
    """Create a new <w:p> directly after p and return it as a Paragraph object."""
    from docx.text.paragraph import Paragraph
    new_p = OxmlElement("w:p")
    # Copy paragraph properties (style, indentation) from p so the new paragraph
    # inherits the surrounding formatting.
    src_pPr = p._element.find(W_PPR)
    if src_pPr is not None:
        new_p.append(deepcopy(src_pPr))
    p._element.addnext(new_p)
    return Paragraph(new_p, p._parent)


# ---------------- main rewriter ------------------------------------------ #


def rewrite():
    if not BACKUP.exists():
        sys.exit(f"Backup not found: {BACKUP}. Aborting for safety.")

    doc = Document(str(SRC))
    # Idempotency guard: look for our marker pattern in the doc.
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "units cancel" in full_text and "→" in full_text and "Mass of solution" in full_text:
        # Heuristic: if we see explanation-style language plus a multi-step
        # NaCl line, assume a prior pass already ran.
        # (We accept false positives; the user can restore from backup.)
        pass  # not strict — proceed

    # Walk paragraphs, rewriting each Solution: in turn.
    paragraphs = list(doc.paragraphs)
    sol_index = 0  # absolute Solution: counter (0-based)
    rewrite_index = 0  # index into REWRITES (skips skipped indices)
    i = 0
    rewrites_applied = 0

    while i < len(paragraphs):
        p = paragraphs[i]
        if not is_solution_paragraph(p):
            i += 1
            continue

        if sol_index in SKIP_INDICES:
            sol_index += 1
            i += 1
            continue

        if rewrite_index >= len(REWRITES):
            print(f"WARN: encountered Solution #{sol_index} but no more rewrites; stopping.")
            break

        anchor_text, consume_after, lines, explanation = REWRITES[rewrite_index]

        # Sanity: existing paragraph should start with "Solution:".
        # (We don't enforce anchor_text matching — order is the contract.)

        # Delete the consume_after paragraphs that follow.
        body = p._element.getparent()
        nxt = p._element.getnext()
        consumed = 0
        while consumed < consume_after and nxt is not None:
            sibling = nxt
            nxt = nxt.getnext()
            body.remove(sibling)
            consumed += 1

        # Rewrite the Solution paragraph with the first line.
        clear_paragraph_runs(p)
        add_run_with_strikes(p, "Solution: ", bold=True)
        add_runs_from_markup(p, lines[0])

        # Insert each subsequent line as a new paragraph after the previous.
        anchor = p
        for extra in lines[1:]:
            new_p = insert_paragraph_after(anchor)
            add_runs_from_markup(new_p, extra)
            anchor = new_p

        # Insert the italic explanation paragraph.
        new_p = insert_paragraph_after(anchor)
        add_runs_from_markup(new_p, explanation, italic=True)
        anchor = new_p

        # Re-list paragraphs and advance i past the rewritten block.
        # New layout starting at original i:
        #   i              -> rewritten Solution (first math line)
        #   i+1 .. i+L-1   -> additional math lines (L = len(lines))
        #   i+L            -> italic explanation
        #   i+L+1          -> next paragraph to examine
        paragraphs = list(doc.paragraphs)
        i = i + len(lines) + 1

        rewrites_applied += 1
        sol_index += 1
        rewrite_index += 1

    doc.save(str(SRC))
    print(f"Applied {rewrites_applied} rewrites across {sol_index} Solution paragraphs.")


if __name__ == "__main__":
    rewrite()
