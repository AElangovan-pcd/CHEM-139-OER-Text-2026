"""Auto-rewrite "Solution:" and "ANSWER:" math chains across every chapter.

Where a Solution paragraph already contains a factor-label expression of the
form

    <number> <unit> [× ( <num> <unit_top> / <num> <unit_bot> )]+ = <result>

this script applies Word strikethrough to every cancelling unit occurrence
and inserts a brief italic explanation paragraph after, naming which units
cancel. Multi-step Solutions whose chain spans a single line are handled.
For Solutions that don't fit the pattern (qualitative answers, formula-style
percent error, IV drip rates with parenthesised denominators, balanced
equations, etc.) the paragraph is left alone -- the script reports the
miss in its run summary so we know which problems still need a hand pass.

The script also handles Worked Example "Step N — ..." paragraphs: any Step
whose body matches the same factor-label pattern receives the same
treatment. The "ANSWER:" line itself is left alone (it is just the result).

Idempotent: a paragraph is skipped if it already contains a <w:strike> run
or if its parent block already contains an italic paragraph that begins
with the auto-explanation prefix "Units cancel as marked".
"""
from copy import deepcopy
from pathlib import Path
import re
import shutil
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(r"C:/Users/easam/Documents/Claude/Projects/OER/CHEM_139_OER_Text_2026")
BACKUP_DIR = ROOT / ".firecrawl" / "backups"

W_T = qn("w:t")
W_R = qn("w:r")
W_RPR = qn("w:rPr")
W_PPR = qn("w:pPr")

# Chapters to process. Chapter 02 was hand-authored already; skip its
# Solutions but still do its Worked Examples.
CHAPTER_FILES = {
    "01": "Chapter_01_Science_and_Measurement.docx",
    "02": "Chapter_02_Unit_Systems_and_Dimensional_Analysis.docx",
    "03": "Chapter_03_Basic_Concepts_About_Matter.docx",
    "04": "Chapter_04_Atoms_Molecules_Subatomic_Particles.docx",
    "05": "Chapter_05_Electronic_Structure_and_Periodicity.docx",
    "06": "Chapter_06_Chemical_Bonds.docx",
    "07": "Chapter_07_Chemical_Nomenclature.docx",
    "08": "Chapter_08_Mole_Concept_and_Chemical_Formulas.docx",
    "09": "Chapter_09_Chemical_Calculations_and_Equations.docx",
    "10": "Chapter_10_States_of_Matter.docx",
}


# ---------------- chain parser ------------------------------------------- #


# A single quantity: number with optional × 10^n exponent, optionally followed
# by a unit phrase. The unit phrase is everything from the first whitespace
# after the number until the next high-level boundary -- including spaces,
# Unicode superscripts (cm³), microns (µ), degree signs (°C), and chemical
# formulas (CO₂, H₂O). Subscript digits ₀-₉ are allowed.
_NUM_RE = r"(?:[-−]?\d+(?:\.\d+)?(?:\s*[×x]\s*10[\d⁻⁰¹²³⁴⁵⁶⁷⁸⁹\-]+)?)"
_QTY_NO_UNIT_RE = re.compile(rf"^\s*({_NUM_RE})\s*$")
_QTY_WITH_UNIT_RE = re.compile(rf"^\s*({_NUM_RE})\s+(.+?)\s*$")
_FACTOR_RE = re.compile(r"^\s*\(\s*(.+?)\s*/\s*(.+?)\s*\)\s*$")


def _parse_qty(s):
    s = s.strip()
    m = _QTY_NO_UNIT_RE.match(s)
    if m:
        return m.group(1), ""
    m = _QTY_WITH_UNIT_RE.match(s)
    if m:
        unit = m.group(2).strip().rstrip(".,;")
        # Reject obvious junk: a unit that contains an operator or digit-only
        # tail is probably a misparse (e.g. expression like "12 + 5").
        if any(ch in unit for ch in "+×=→"):
            return None
        return m.group(1), unit
    return None


def _split_top_level(s, sep="×"):
    """Split s on `sep` characters that are NOT inside parentheses."""
    parts = []
    depth = 0
    start = 0
    i = 0
    while i < len(s):
        c = s[i]
        if c == "(":
            depth += 1
        elif c == ")":
            depth -= 1
        elif c == sep and depth == 0:
            parts.append(s[start:i])
            start = i + 1
        i += 1
    parts.append(s[start:])
    return parts


def parse_chain(text):
    """Parse `<initial> [× <factor>]* = <result>` into a structured form.

    Returns dict with 'initial', 'factors', 'result' OR None if not parseable.
    """
    # Trim trailing punctuation and surrounding whitespace.
    text = text.strip()
    text = re.sub(r"[.;]+$", "", text).strip()

    # Find the first top-level "=" (or the chain's only "="). Some Solutions
    # have multiple equalities like "= 8.00 cm³ = 8.00 mL"; we keep the
    # whole RHS verbatim and only parse the LHS.
    depth = 0
    eq_pos = -1
    for i, c in enumerate(text):
        if c == "(":
            depth += 1
        elif c == ")":
            depth -= 1
        elif c == "=" and depth == 0:
            eq_pos = i
            break
    if eq_pos < 0:
        return None
    lhs = text[:eq_pos].strip()
    rhs = text[eq_pos + 1:].strip()

    factors_raw = _split_top_level(lhs, "×")
    if len(factors_raw) < 2:
        # No × means no factor-label form.
        return None

    initial = _parse_qty(factors_raw[0])
    if initial is None:
        return None

    parsed = []
    for raw in factors_raw[1:]:
        m = _FACTOR_RE.match(raw)
        if m is None:
            return None
        top = _parse_qty(m.group(1))
        bot = _parse_qty(m.group(2))
        if top is None or bot is None:
            return None
        if not top[1] and not bot[1]:
            # Pure number ratio with no units: doesn't help cancellation logic.
            return None
        parsed.append({"top": top, "bot": bot, "raw_top": m.group(1).strip(), "raw_bot": m.group(2).strip()})

    return {"initial": initial, "factors": parsed, "result": rhs, "lhs_raw": lhs}


# ---------------- cancellation analysis ---------------------------------- #


def find_cancellers(parsed):
    """Determine which units cancel.

    A unit token cancels when the same unit appears at least once in a
    numerator position AND at least once in a denominator position on the
    LHS of the chain (excluding the result).
    """
    nums = []
    if parsed["initial"][1]:
        nums.append(parsed["initial"][1])
    for f in parsed["factors"]:
        if f["top"][1]:
            nums.append(f["top"][1])
    dens = [f["bot"][1] for f in parsed["factors"] if f["bot"][1]]

    nums_set = set(nums)
    dens_set = set(dens)
    return nums_set & dens_set


# ---------------- text reconstruction ------------------------------------ #


def rebuild_chain(parsed, cancellers):
    """Return the reconstructed math line text and a list of token-segments
    suitable for emitting to a paragraph (with strikethrough on cancellers).

    Segments are tuples of (text, struck_bool).
    """
    segs = []

    def emit(text, struck=False):
        segs.append((text, struck))

    # Initial: number + unit
    init_num, init_unit = parsed["initial"]
    emit(f"{init_num} ")
    if init_unit:
        if init_unit in cancellers:
            emit(init_unit, True)
        else:
            emit(init_unit)

    for f in parsed["factors"]:
        emit(" × (")
        emit(f["top"][0])
        if f["top"][1]:
            emit(" ")
            if f["top"][1] in cancellers:
                emit(f["top"][1], True)
            else:
                emit(f["top"][1])
        emit(" / ")
        emit(f["bot"][0])
        if f["bot"][1]:
            emit(" ")
            if f["bot"][1] in cancellers:
                emit(f["bot"][1], True)
            else:
                emit(f["bot"][1])
        emit(")")

    emit(" = ")
    emit(parsed["result"])

    return segs


def make_explanation(cancellers, parsed):
    units_quoted = ", ".join('"' + u + '"' for u in sorted(cancellers))
    if parsed["factors"]:
        # Identify the surviving (un-cancelled) numerator unit, if any.
        survivors = []
        if parsed["initial"][1] and parsed["initial"][1] not in cancellers:
            survivors.append(parsed["initial"][1])
        for f in parsed["factors"]:
            if f["top"][1] and f["top"][1] not in cancellers:
                survivors.append(f["top"][1])
        survivor_text = ""
        if survivors:
            tail_unit = survivors[-1]
            survivor_text = f' The remaining unit is "{tail_unit}".'
    else:
        survivor_text = ""
    return (
        f"Units cancel as marked: {units_quoted} cancel(s) where the "
        f"strikethrough crosses them out.{survivor_text}"
    )


# ---------------- docx run surgery --------------------------------------- #


def _add_run(p, text, *, bold=False, italic=False, struck=False):
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if struck:
        rPr = run._element.get_or_add_rPr()
        s_el = OxmlElement("w:strike")
        s_el.set(qn("w:val"), "true")
        rPr.append(s_el)
    return run


def _clear_runs(p):
    p_el = p._element
    for child in list(p_el):
        if child.tag != W_PPR:
            p_el.remove(child)


def _insert_paragraph_after(anchor):
    new_p = OxmlElement("w:p")
    src_pPr = anchor._element.find(W_PPR)
    if src_pPr is not None:
        new_p.append(deepcopy(src_pPr))
    anchor._element.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def already_struck(p):
    return "w:strike" in p._element.xml


def already_explained(p, doc_paragraphs, idx):
    """Look ahead two paragraphs for an italic auto-explanation marker."""
    for j in range(idx + 1, min(idx + 4, len(doc_paragraphs))):
        text = doc_paragraphs[j].text or ""
        if text.startswith("Units cancel as marked"):
            return True
        if text.strip().startswith("Solution:") or text.strip().startswith("ANSWER:"):
            break
    return False


# ---------------- main per-chapter rewriter ------------------------------ #


_LABEL_PATTERNS = [
    # The Solution body sits after one of these labels at the start of a paragraph.
    (re.compile(r"^Solution:\s*", re.IGNORECASE), "Solution: "),
    (re.compile(r"^Step\s+\d+\s*[—–-]\s*", re.IGNORECASE), None),  # keep original prefix
]


def _split_label(text):
    """Return (prefix_text, body_text) if the paragraph starts with a known
    rewrite-eligible label. Else (None, None)."""
    for pat, replacement in _LABEL_PATTERNS:
        m = pat.match(text)
        if m:
            prefix = replacement if replacement is not None else m.group(0)
            body = text[m.end():]
            return prefix, body
    return None, None


def process_chapter(num, filename, *, skip_solutions_for_ch="02"):
    src = ROOT / filename
    if not src.exists():
        print(f"Ch {num}: missing {src}")
        return 0, 0, 0

    doc = Document(str(src))
    paragraphs = list(doc.paragraphs)

    rewritten = 0
    skipped_already = 0
    skipped_unparseable = 0

    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        text = p.text or ""
        prefix, body = _split_label(text)
        if prefix is None:
            i += 1
            continue

        # Chapter-skip: don't touch hand-authored Chapter 2 Solutions, but DO
        # touch its Worked Example "Step N —" paragraphs.
        is_solution = prefix.startswith("Solution")
        if num == skip_solutions_for_ch and is_solution:
            i += 1
            continue

        # Idempotency.
        if already_struck(p) or already_explained(p, paragraphs, i):
            skipped_already += 1
            i += 1
            continue

        # Try to parse the body as a math chain.
        parsed = parse_chain(body)
        if parsed is None:
            skipped_unparseable += 1
            i += 1
            continue

        cancellers = find_cancellers(parsed)
        if not cancellers:
            skipped_unparseable += 1
            i += 1
            continue

        # Rebuild paragraph runs.
        segs = rebuild_chain(parsed, cancellers)
        _clear_runs(p)
        _add_run(p, prefix, bold=is_solution)
        for txt, struck in segs:
            _add_run(p, txt, struck=struck)

        # Insert italic explanation paragraph immediately after.
        explanation_p = _insert_paragraph_after(p)
        _add_run(explanation_p, make_explanation(cancellers, parsed), italic=True)

        # Re-list paragraphs because we mutated the body.
        paragraphs = list(doc.paragraphs)
        i = i + 2  # skip past explanation
        rewritten += 1

    if rewritten:
        doc.save(str(src))
    return rewritten, skipped_already, skipped_unparseable


def main():
    print("Auto-rewriter for Solution / Step math chains.\n")
    totals = {"rewritten": 0, "already": 0, "unparseable": 0}
    for num, filename in CHAPTER_FILES.items():
        r, a, u = process_chapter(num, filename)
        totals["rewritten"] += r
        totals["already"] += a
        totals["unparseable"] += u
        print(f"Ch {num}: rewrote {r:>3}, skipped {a} already-done, {u} unparseable.")
    print()
    print(
        f"TOTAL: {totals['rewritten']} rewritten, "
        f"{totals['already']} already, "
        f"{totals['unparseable']} left for hand pass."
    )


if __name__ == "__main__":
    main()
